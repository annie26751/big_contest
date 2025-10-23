# app.py

import streamlit as st
import pandas as pd
import numpy as np
import os
import io
from docx import Document
import time
from data_processor import load_fixed_data, analyze_merchant, FIXED_DATA_PATH
from gemini_api import generate_marketing_text_with_gemini, generate_chat_response_with_gemini
from visualize import load_data
from visualize import kpi_board, gender_age_pie, customer_type_pie_revisit_new, customer_type_pie_origin
from mbti_classifier import classify_merchant_mbti
from visualization_area import render_area_dashboard
from clustering import get_dtw_cluster, build_dtw_report

def load_area_cluster_data():
    """area_cluster.csv 파일을 로드하는 캐시 함수"""
    # data_dong.csv가 있는 폴더(FIXED_DATA_PATH) 기준으로 area_cluster.csv 경로 설정
    cluster_file_path = os.path.join(os.path.dirname(FIXED_DATA_PATH), "area_clustering.csv")
    
    if not os.path.exists(cluster_file_path):
        st.error(f"❌ 'area_cluster.csv' 파일을 찾을 수 없습니다. 경로: {cluster_file_path}")
        return None
    try:
        # area_cluster.csv 파일 인코딩에 맞게 'utf-8' 또는 'cp949' 시도
        df = pd.read_csv(cluster_file_path, encoding='utf-8')
    except UnicodeDecodeError:
        df = pd.read_csv(cluster_file_path, encoding='cp949')
    except Exception as e:
        st.error(f"❌ 'area_cluster.csv' 로드 중 오류: {e}")
        return None
    return df

@st.cache_resource(ttl=3600)
def cached_load_data(path):
    """Streamlit 캐싱을 적용한 데이터 로드 함수"""
    try:
        return load_fixed_data(path)
    except (FileNotFoundError, UnicodeDecodeError, ValueError, Exception) as e:
        st.error(f"데이터 로드 중 오류 발생: {e}")
        st.info("FIXED_DATA_PATH 변수가 정확한지, 파일이 해당 경로에 존재하는지 확인해주세요.")
        st.stop()
        return None
    
def create_docx_report(mct_id, proposal, chat_history):
    """마케팅 전략과 챗봇 대화 내용으로 Word 문서를 생성하여 바이트 객체로 반환"""
    doc = Document()
    doc.add_heading(f"'{mct_id}' 가맹점 AI 마케팅 분석 리포트", level=1)
    doc.add_paragraph()

    doc.add_heading("🚀 AI 비밀상담사의 맞춤형 마케팅 플랜", level=2)
    for line in proposal.split('\n'):
        doc.add_paragraph(line)
    
    if len(chat_history) > 1:
        doc.add_paragraph()
        doc.add_heading("🤖 추가 상담 내용 (Q&A)", level=2)
        for message in chat_history[1:]:
            role = "Q (사용자)" if message["role"] == "user" else "A (AI 상담사)"
            p = doc.add_paragraph()
            p.add_run(f"{role}: ").bold = True
            p.add_run(message['content'])
            doc.add_paragraph()
            
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def create_cluster_report_docx(
    mct_id: str,
    h_dong: str,
    industry_name: str,
    selected_industry_mapped: str,
    cluster_text: str,
    similar_dong_sentence: str,
    cluster_description: str,
    all_desc: str
) -> bytes:
    """업장 보고서(상권 클러스터) 탭의 내용을 Word 문서로 생성하여 바이트 객체로 반환"""
    doc = Document()
    doc.add_heading(f"'{mct_id}' 가맹점 상권(업장) 분석 리포트", level=1)
    doc.add_paragraph()

    # 1. 핵심 분석 결과
    doc.add_heading("📌 핵심 분석 결과", level=2)
    
    # h_dong, industry_name 등 주요 정보를 굵게 처리
    p1 = doc.add_paragraph()
    p1.add_run(f"점주님의 업장 (").bold = False
    p1.add_run(f"{h_dong}, {industry_name}").bold = True
    p1.add_run(f")은(는) ").bold = False
    p1.add_run(f"[{selected_industry_mapped}-{cluster_text}]").bold = True
    p1.add_run("에 해당합니다.").bold = False
    
    # 유사 행정동 문장 (굵게 처리된 마크다운 제거)
    if similar_dong_sentence:
        clean_sentence = similar_dong_sentence.replace("**", "") # 마크다운 ** 제거
        clean_sentence = clean_sentence.replace("[", "").replace("]", "") # [ ] 제거
        doc.add_paragraph(clean_sentence)
        
    doc.add_paragraph() # 여백

    # 2. 상세 분석 (해당 클러스터)
    if cluster_description:
        doc.add_heading(f"➡️ {cluster_text} 상세 분석", level=2)
        # 텍스트 파일의 줄바꿈(개행)을 docx에 반영
        for line in cluster_description.split('\n'):
            doc.add_paragraph(line)
        doc.add_paragraph() # 여백

    # 3. 업종 전체 요약
    if all_desc:
        doc.add_heading(f"'{selected_industry_mapped}' 업종 전체 요약", level=2)
        for line in all_desc.split('\n'):
            doc.add_paragraph(line)
            
    # 버퍼에 저장하여 반환
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# -------------------- 메인 시작 -------------------- #
def main():
    st.set_page_config(layout="wide", page_title="💡 내 가게를 살리는 AI 비밀상담사")
    st.markdown("""
    <div style="background-color:#f0f2f6; padding: 1rem; border-radius: 10px; margin-bottom: 1rem;">
    <h1 style="text-align: center; color: black; margin: 0; font-size: 2.5rem;">💡 내 가게를 살리는 AI 비밀상담사</h1>
    </div>
    """, unsafe_allow_html=True)
    #st.markdown("데이터와 AI를 기반으로 우리 가게의 문제를 진단하고, 핵심 고객을 위한 맞춤 마케팅 전략을 찾아보세요.")

# --- Session State 초기화 --- #
    if 'generating' not in st.session_state:
        st.session_state['generating'] = False
    if 'chat_messages' not in st.session_state:
        st.session_state['chat_messages'] = []
    
# --- 데이터 로드 & 캐시 --- #
    if 'df_profile' not in st.session_state:
        with st.spinner('초기 데이터를 로드하는 중입니다...'):
            df_profile = cached_load_data(FIXED_DATA_PATH)
            if df_profile is None:
                st.error("데이터 로드에 실패하여 앱을 실행할 수 없습니다.")
                st.stop()
            st.session_state['df_profile'] = df_profile

    df_profile = st.session_state['df_profile']

    # --- 가맹점 ID 캐시 --- #
    if 'merchant_ids' not in st.session_state:
        merchant_ids = (
            df_profile['ENCODED_MCT']
            .dropna()
            .astype(str)
            .unique()
            .tolist()
        )
        merchant_ids.sort()
        st.session_state['merchant_ids'] = merchant_ids

    all_merchant_ids = st.session_state['merchant_ids']

    st.sidebar.header("시작하기")

    # -------------------- 사이드바 -------------------- #
    # 1. 가맹점 선택
    merchant_ids = df_profile["ENCODED_MCT"].unique().tolist()

    with st.sidebar:
        search_term = st.text_input(
            "가맹점 선택하기",
            placeholder="여기에 가맹점 번호를 입력하세요"
        ).strip()

        if search_term:
            filtered_merchants = [m for m in merchant_ids if search_term in m]
        else:
            filtered_merchants = merchant_ids

        # 검색 결과가 있으면 selectbox 표시, 없으면 안내
        if filtered_merchants:
            selected_mct = st.selectbox(
                "분석할 가맹점을 선택하세요:",
                filtered_merchants,
                key="merchant_selector",
                label_visibility="collapsed"
            )
        else:
            st.info("일치하는 가맹점이 없습니다.")
            selected_mct = None

    if not selected_mct:
        st.info("사이드바에서 분석할 가맹점을 선택해주세요.")
        st.stop()


    # -------------------- 가맹점 기본 정보 블록 -------------------- #
    if selected_mct:
        try:
            row = df_profile.loc[df_profile["ENCODED_MCT"].astype(str) == str(selected_mct)].iloc[0]
        except IndexError:
            st.sidebar.info("선택한 가맹점의 기본 정보를 찾을 수 없습니다.")
        else:
            with st.sidebar.expander("📂 가게 정보", expanded=True):
                status = "운영 중" if pd.isna(row.get("MCT_ME_D")) else f"폐업 ({row.get('MCT_ME_D')})"
                st.markdown(f"**업종:** {row.get('HPSN_MCT_ZCD_NM')}")
                st.markdown(f"**주소:** {row.get('MCT_BSE_AR')}")
                st.markdown(f"**상권:** {row.get('h_name', row.get('HPSN_MCT_BZN_CD_NM'))}")
                st.markdown(f"**상태:** {status}")

            # MBTI(가게 유형) 분류 — 기존 UI 유지
            store_type = classify_merchant_mbti(row)
            with st.sidebar.expander("🏪 가게 유형 (MBTI)", expanded=True):
                st.markdown(f"**{store_type['name']}**")
                st.caption(store_type['description'])
    else:
        st.sidebar.info("가맹점을 선택하면 기본 정보와 가게 유형이 표시됩니다.")

    # -------------------- 기준월 선택 -------------------- #
    st.session_state["selected_mct"] = selected_mct
    @st.cache_data(ttl=3600)
    def _load_viz_df():
        return load_data()

    if "viz_df" not in st.session_state:
        st.session_state["viz_df"] = _load_viz_df()
    viz_df = st.session_state["viz_df"]

    col_title, col_month = st.columns([3, 1], gap="large")
    with col_title:
        pass  # 이미 위에서 타이틀/설명 렌더링

    with col_month:
        mct_for_month = st.session_state.get("selected_mct")  # ← 지역변수 대신 세션값 사용
        if mct_for_month:
            mct_dates = (
                viz_df.loc[viz_df["ENCODED_MCT"] == mct_for_month, "TA_YM_DT"]
                .dropna().sort_values(ascending=False)
                .dt.to_period("M").drop_duplicates().tolist()
            )
            month_options = ["-- 기준월을 선택하세요 --"] + [p.to_timestamp() for p in mct_dates]
        else:
            month_options = ["-- 기준월을 선택하세요 --"]

        selected_ref = st.selectbox(
            "📅 날짜",
            month_options, index=0,
            format_func=lambda d: d if isinstance(d, str) else d.strftime("%Y-%m"),
            key="ref_month_selector",
        )
        if selected_ref == "-- 기준월을 선택하세요 --":
            selected_ref = None


    # -------------------- 메인 화면 -------------------- #
    # 1️⃣ 가맹점 선택 후 분석 실행
    if selected_mct is None:
        st.session_state["analysis_result"] = None
    else:
        df_mct = df_profile[df_profile["ENCODED_MCT"] == selected_mct]
        if df_mct.empty:
            st.warning("선택한 가맹점 데이터가 없습니다.")
            st.session_state["analysis_result"] = None
        else:
            if (
                "last_mct" not in st.session_state
                or "last_ref" not in st.session_state
                or st.session_state["last_mct"] != selected_mct
                or st.session_state["last_ref"] != selected_ref
            ):
                with st.spinner("가맹점 데이터 분석 중..."):
                    st.session_state["analysis_result"] = analyze_merchant(df_mct.iloc[0])

                st.session_state["last_mct"] = selected_mct
                st.session_state["last_ref"] = selected_ref
                st.session_state["marketing_proposal"] = ""
                st.session_state["show_mbti_description"] = False
    # 2️⃣ 탭 선언
    tab_llm, tab_viz, tab_area, tab_clu = st.tabs(["🤖 AI 마케팅", "📊 월별 보고서", "📍 상권 보고서", "❤️ 업장 보고서"])

    # 3️⃣ 안전 가드
    if (
        "analysis_result" not in st.session_state
        or st.session_state["analysis_result"] is None
    ):
        st.warning("아직 가맹점을 선택하지 않았습니다. 사이드바에서 먼저 선택하세요.")
        st.stop()

    # 4️⃣ 세션에서 결과 꺼내오기
    analysis_result = st.session_state["analysis_result"]
    summary = analysis_result["summary"]
    persona = analysis_result["persona"]
    mbti_result = analysis_result["mbti"]
    mct_data = analysis_result["raw_data"]



    with tab_clu:
        st.header("❤️ 1. 상권 그룹 분석")
        # 1. area_cluster.csv에서 사용할 업종명 매핑
        INDUSTRY_MAP = {
            "🍗 고기집 🍗": "고기집",
            "🍎 식료품 🍎": "식료품",
            "🥐 베이커리 🥐": "베이커리",
            "🍺 술집 🍺": "술집",
            "🍝 양식음식점 🍝": "양식음식점",
            "🍣 일식음식점 🍣": "일식음식점",
            "👲 중식음식점 👲": "중식음식점",
            "🍔 패스트푸드점 🍔": "패스트푸드점",
            "🍚 한식음식점 🍚": "한식음식점"
        }
        
        # 2. 클러스터 데이터 로드
        df_area_cluster = load_area_cluster_data()
        if df_area_cluster is None:
            st.error("'area_cluster.csv' 로드에 실패하여 업장 보고서를 표시할 수 없습니다.")
            st.stop()

        # 3. 가맹점 기본 정보(행정동, 업종명) 추출
        try:
            mct_row = df_profile.loc[df_profile["ENCODED_MCT"] == selected_mct].iloc[0]
            h_name = mct_row.get("h_name", "정보없음")
            h_dong = h_name.split(' ')[-1] # 행정동
            industry_name = mct_row.get("HPSN_MCT_ZCD_NM", "정보없음") # 업종

        except IndexError:
            st.error("선택한 가맹점의 기본 정보(h_name, 업종)를 찾을 수 없습니다.")
            st.stop()
        except Exception as e:
            st.error(f"가맹점 정보 처리 중 오류: {e}")
            st.stop()
            
        # 4. 사용자에게 비교 업종 선택받기
        st.subheader(f"점주님의 업장 ({industry_name})이 위치한 [{h_dong}]의 상권 특성을 비교 분석합니다.")
        
        # 4-1. 선택 옵션 리스트에 안내 문구 추가
        industry_options = ["-- 점주님의 업종을 선택하세요 --"] + list(INDUSTRY_MAP.keys())

        selected_industry_emoji = st.selectbox(
            "👇 업종을 선택하면, 아래에 상권 분석 결과가 표시됩니다 ",
            industry_options,
            index=0, # 기본값으로 0번째(-- 선택하세요 --)를 지정
            key="area_cluster_industry_select"
        )

        st.markdown("---")
        
        # --- ⬇️ 업종이 선택되었을 때만 하단 내용 표시 ---
        if selected_industry_emoji != "-- 점주님의 업종을 선택하세요 --":
            
            # 4-2. 선택된 경우에만 매핑 수행
            selected_industry_mapped = INDUSTRY_MAP[selected_industry_emoji]

            # 5. 클러스터 번호 및 유사 행정동 찾기 (들여쓰기)
            cluster_text = "클러스터 정보 없음"
            similar_dong_sentence = "" 
            cluster_num = None 
            cluster_description = "" # DOCX용 변수 초기화
            all_desc = ""            # DOCX용 변수 초기화
            
            try:
                cluster_row = df_area_cluster[
                    (df_area_cluster['서비스_업종_코드_명'] == selected_industry_mapped) &
                    (df_area_cluster['행정동_코드_명'] == h_dong)
                ]

                if not cluster_row.empty:
                    cluster_num = cluster_row['area_cluster'].iloc[0] 
                    cluster_text = f"Cluster {cluster_num}"

                    # ... (유사 행정동 찾기 로직) ...
                    similar_dongs_df = df_area_cluster[
                        (df_area_cluster['area_cluster'] == cluster_num) &
                        (df_area_cluster['서비스_업종_코드_명'] == selected_industry_mapped)
                    ]
                    similar_dongs_list = similar_dongs_df['행정동_코드_명'].unique().tolist()
                    similar_dongs_list = [dong for dong in similar_dongs_list if dong != h_dong]
                    
                    if similar_dongs_list:
                        similar_dongs_str = ", ".join(similar_dongs_list)
                        similar_dong_sentence = f"[{h_dong}]과 유사한 추이를 보이는 행정동으로는 [{similar_dongs_str}]이 있습니다."
                    else:
                        similar_dong_sentence = f"[{h_dong}]과 유사한 추이를 보이는 다른 행정동이 없습니다."

                else:
                    st.warning(f"참고: '{h_dong}'의 '{selected_industry_mapped}' 업종에 대한 클러스터 정보가 'area_cluster.csv'에 없습니다.")

            except KeyError as e:
                st.error(f"'area_cluster.csv' 파일에 필요한 컬럼({e})이 없습니다. (서비스_업종_코드_명, 행정동_코드_명, area_cluster)")
                cluster_text = "오류"
            except Exception as e:
                st.error(f"클러스터 조회 중 오류: {e}")
                cluster_text = "오류"

            # 6. 최종 결과 문구 표시
            st.markdown("---")
            st.markdown(f"점주님의 업장은 [{h_dong}]에 위치한 [{industry_name}] 이며, **[{selected_industry_mapped}-{cluster_text}]**에 해당합니다.")
            
            # 7. 유사 행정동 문장 출력
            if similar_dong_sentence: 
                st.markdown(similar_dong_sentence)
            
            # 8-1. 업종 전체 클러스터 요약 (cluster.txt) 및 이미지 표시
            st.markdown(f"### 점주님의 업장이 속한 **[{selected_industry_mapped} - {cluster_text}]**의 특징을 알아볼까요?😊")
            

            cluster_summary_path = f"./text/{selected_industry_mapped}/cluster.txt"
            if os.path.exists(cluster_summary_path):
                try:
                    with open(cluster_summary_path, 'r', encoding='utf-8') as f:
                        cluster_summary_desc = f.read()
                except UnicodeDecodeError:
                    with open(cluster_summary_path, 'r', encoding='cp949') as f:
                        cluster_summary_desc = f.read()
                except Exception as e:
                    cluster_summary_desc = f"파일 로드 오류: {e}"
                
                # st.subheader(f"'{selected_industry_mapped}' 업종 클러스터 요약") # 제목이 필요하면 주석 해제
                st.markdown(cluster_summary_desc) # 요약 텍스트 표시
            else:
                 st.caption(f"[클러스터 요약 없음: {cluster_summary_path}]")
                
                
            image_path = f"./image/{selected_industry_mapped}.png"
            if os.path.exists(image_path):
                st.image(image_path, caption=f"'{selected_industry_mapped}' 업종 클러스터 분포")
            else:
                st.caption(f"[이미지 없음: {image_path}]") 

            # 8-2. 특정 클러스터 분석 텍스트 표시
            if cluster_num is not None: 
                text_path = f"./text/{selected_industry_mapped}/cluster{cluster_num}.txt"
                if os.path.exists(text_path):
                    try:
                        with open(text_path, 'r', encoding='utf-8') as f:
                            cluster_description = f.read() # 변수에 저장
                    except UnicodeDecodeError:
                        with open(text_path, 'r', encoding='cp949') as f:
                            cluster_description = f.read() # 변수에 저장
                    except Exception as e:
                        cluster_description = f"텍스트 파일 로드 오류: {e}"
                    
                    st.subheader(f"➡️ {cluster_text} 상세 분석")
                    st.markdown(cluster_description)
                else:
                    st.caption(f"[분석 내용 없음: {text_path}]")
                
                # ... (다른 클러스터 토글 로직) ...
                with st.expander(f"다른 '{selected_industry_mapped}' 클러스터 유형 살펴보기"):
                    all_cluster_nums = sorted(df_area_cluster[
                        df_area_cluster['서비스_업종_코드_명'] == selected_industry_mapped
                    ]['area_cluster'].unique())
                    
                    found_other = False
                    for c_num in all_cluster_nums:
                        if c_num == cluster_num: continue
                        found_other = True
                        other_text_path = f"./text/{selected_industry_mapped}/cluster{c_num}.txt"
                        if os.path.exists(other_text_path):
                            try:
                                with open(other_text_path, 'r', encoding='utf-8') as f: other_desc = f.read()
                            except UnicodeDecodeError:
                                with open(other_text_path, 'r', encoding='cp949') as f: other_desc = f.read()
                            except Exception as e: other_desc = f"파일 로드 오류: {e}"
                            st.markdown("---"); st.subheader(f"Cluster {c_num} 분석"); st.markdown(other_desc)
                        else:
                            st.caption(f"[분석 내용 없음: {other_text_path}]")
                    if not found_other: st.caption("다른 클러스터 정보가 없습니다.")
            else: 
                st.caption(f"[클러스터 정보 없음: '{h_dong}'의 '{selected_industry_mapped}' 데이터 확인 필요]")

            # 8-3. 업종별 all.txt 파일 불러오기
            st.markdown("---")
            all_text_path = f"./text/{selected_industry_mapped}/all.txt"
            if os.path.exists(all_text_path):
                st.subheader(f"'{selected_industry_mapped}' 업종 전체 요약")
                try:
                    with open(all_text_path, 'r', encoding='utf-8') as f:
                        all_desc = f.read() # 변수에 저장
                except UnicodeDecodeError:
                    with open(all_text_path, 'r', encoding='cp949') as f:
                        all_desc = f.read() # 변수에 저장
                except Exception as e:
                    all_desc = f"파일 로드 오류: {e}"
                
                st.markdown(all_desc)
            else:
                st.caption(f"[전체 요약 없음: {all_text_path}]")

            # 9. DOCX 다운로드 버튼
            
            # DOCX 데이터 생성
            docx_data_clu = create_cluster_report_docx(
                mct_id=selected_mct,
                h_dong=h_dong,
                industry_name=industry_name,
                selected_industry_mapped=selected_industry_mapped,
                cluster_text=cluster_text,
                similar_dong_sentence=similar_dong_sentence,
                cluster_description=cluster_description,
                all_desc=all_desc
            )
            
            st.download_button(
                label="📄 상권(업장) 분석 내용 저장하기 (.docx)",
                data=docx_data_clu,
                file_name=f"report_area_{selected_mct}_{selected_industry_mapped}.docx", 
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="docx_download_clu"
            )
            st.markdown("---") # 구분선
        
        else:
            # 4-3. 업종이 선택되지 않았을 때
            st.markdown("")

        st.header("❤️ 2. 우리 가게는 어떤 유형일까요?")
        # 선택된 가맹점 이름 가져오기 (없으면 ID로 대체)
        mct_row = df_profile.loc[df_profile["ENCODED_MCT"] == selected_mct]
        merchant_name = (
            mct_row["MCT_NM"].iloc[0]
            if ("MCT_NM" in mct_row.columns and not mct_row.empty and pd.notna(mct_row["MCT_NM"].iloc[0]))
            else str(selected_mct)
        )

        # 클러스터 찾기
        cluster_id = get_dtw_cluster(selected_mct)
        if cluster_id is None:
            st.info("이 가맹점은 아직 DTW 군집이 매핑되어 있지 않음..")
            st.stop()

        # 리포트 생성 (clustering.py에서 문구/이미지 경로 구성)
        report = build_dtw_report(selected_mct, merchant_name)

        # 헤더/인트로
        st.subheader(report["intro_title"])
        st.markdown(
    f"<p style='font-size:1.1rem; line-height:1.6; font-weight:500;'>{report['intro_body']}</p>",
    unsafe_allow_html=True
)

        # 클러스터 뱃지 느낌
        badge = report["cluster_badge"]
        st.markdown(f"{badge['icon']} **{badge['name']}**", unsafe_allow_html=True)

        # 패턴/해석/주요 업종
        with st.container(border=True):
            st.write(report["pattern"])
            st.write(report["key_industries"])
            st.write(report["interpretation"])
            

        # 그래프 이미지
        imgs = report.get("images", [])
        if imgs:
            st.markdown("**같은 그룹의 상위 10개 업종이에요**")
            for p in imgs:
                st.image(p, caption=os.path.basename(p), use_container_width=True)
        else:
            st.caption("이미지가 없어요. (data/plots 경로/파일명을 확인하세요)")

        # 메타
        st.caption(f"모델: {report['meta']['model_ver']} · 소스: {report['meta']['data_source']}")


    with tab_area:
        @st.cache_data(ttl=3600, show_spinner=False)
        def _auto_load_df_filtered():
            import os
            base_dir = os.path.dirname(FIXED_DATA_PATH)
            csv_path = os.path.join(base_dir, "mapping.csv")
            if not os.path.exists(csv_path):
                raise FileNotFoundError(f"mapping.csv를 찾을 수 없습니다: {csv_path}")
            try:
                df = pd.read_csv(csv_path, encoding="utf-8-sig")
            except UnicodeDecodeError:
                df = pd.read_csv(csv_path, encoding="utf-8")
            return df, base_dir

        try:
            df_filtered, base_dir = _auto_load_df_filtered()
            st.caption(f"🔄 자동 로드: mapping.csv — {len(df_filtered):,}행")
        except Exception as e:
            st.error(f"자동 로드 중 오류: {e}")
            st.stop()

        # 선택 가맹점(ENCODED_MCT)과 base_dir을 넘겨 자동 업종 매핑 적용
        render_area_dashboard(df_filtered, selected_mct=selected_mct, base_dir=base_dir)


    with tab_viz:

        # KPI 비교 차트
        df = load_data()
        st.subheader("📊 저번 달과 이만큼 달라요")
        kpi_board(df, selected_mct, REF=selected_ref)
        st.markdown("---")
        st.subheader("👥 우리 가게를 찾은 손님들")
        st.write("") 

        col1, col2, col3 = st.columns([1, 1, 1], gap="medium")
        with col1:
            gender_age_pie(df, selected_mct, REF=selected_ref)  # 성별·연령 도넛

        with col2:
            customer_type_pie_revisit_new(df, selected_mct, REF=selected_ref)  # 재방문 vs 신규

        with col3:
            customer_type_pie_origin(df, selected_mct, REF=selected_ref) # 거주/직장/유동

    with tab_llm:

        # 1. 데이터 기반 핵심 진단
        st.subheader("🤖 AI가 확인한 우리 가게의 현재 상태는?")
        st.success(f"**[고객층 분석]** {summary['cust_analysis_text']}")
        st.info(f"**[고객 유지력]** {summary['retention_analysis_text']}")
        st.warning(f"**[경쟁 환경]** {summary['comp_analysis_text']}")
        st.markdown("---")
        
        # 2. 페르소나 분석 결과
        st.subheader("🎯 우리 가게의 핵심 고객은 누구일까요?")
        
        with st.container(border=True):
            persona_icon = persona.get("icon", "")
            st.markdown(f"#### {persona_icon} {persona['name']}")

            description_html = persona['description'].replace('\n', '<br>')
            goals_html = "<ul>" + "".join([f"<li>{g}</li>" for g in persona['goals']]) + "</ul>"
            pain_points_html = "<ul>" + "".join([f"<li>{p}</li>" for p in persona['pain_points']]) + "</ul>"

            st.markdown(f"""
            <style>
                .persona-table {{
                    width: 100%;
                    border-collapse: collapse;
                    border: none;
                }}
                .persona-table th, .persona-table td {{
                    border: 1px solid #dfe3e8;
                    padding: 12px 15px;
                    text-align: left;
                    vertical-align: top;
                }}
                .persona-table th {{
                    background-color: #f0f2f6;
                    width: 20%;
                    font-weight: bold;
                }}
                .persona-table ul {{
                    padding-left: 20px;
                    margin: 0;
                }}
            </style>
            <table class="persona-table">
                <tr>
                    <th>소개</th>
                    <td>{description_html}</td>
                </tr>
                <tr>
                    <th>찾는 이유 <br>(Goals)</th>
                    <td>{goals_html}</td>
                </tr>
                <tr>
                    <th>겪는 어려움 <br>(Pain Points)</th>
                    <td>{pain_points_html}</td>
                </tr>
            </table>
            """, unsafe_allow_html=True)
        
        st.markdown("---")

        st.subheader("🧠 AI 비밀상담사의 맞춤형 마케팅 플랜")
        st.markdown("아래는 입력된 데이터와 페르소나를 기반으로 Gemini AI 마케팅 전략을 실시간으로 생성합니다.")
        
        with st.expander("🎯 (선택) 타겟 페르소나 직접 설정하기", expanded=False):
            st.info("특정 고객층을 대상으로 전략을 생성하고 싶다면, 아래에서 직접 페르소나(타겟)을 설정하세요.")
            
            c1, c2, c3 = st.columns(3)
            with c1:
                target_gender = st.selectbox(
                    "성별", 
                    ["데이터 기반", "남성", "여성","남성 및 여성"], 
                    key="target_gender_select"
                )
            with c2:
                # 사용자가 요청한 '10대'를 '10-20대'로 통합
                target_age = st.selectbox(
                    "나이", 
                    ["데이터 기반", "10-20대", "30대", "40대", "50대", "60대 이상"], 
                    key="target_age_select"
                )
            with c3:
                # 사용자가 요청한 '가족'을 데이터의 '거주'와 매핑
                target_cust_type = st.selectbox(
                    "고객 유형", 
                    ["데이터 기반", "직장인", "유동인구", "가족/거주"], 
                    key="target_type_select"
                )
                
        _, btn_col, _ = st.columns([1, 2, 1])
        with btn_col:
            button_text = "🚀 생성중..." if st.session_state.generating else "🚀 AI 마케팅 전략 생성하기"
            if st.button(button_text, use_container_width=True, type="primary", disabled=st.session_state.generating):
                st.session_state.generating = True
                
                override_target = {}
                if target_gender != "데이터 기반":
                    override_target['gender'] = target_gender
                if target_age != "데이터 기반":
                    override_target['age'] = target_age
                if target_cust_type != "데이터 기반":
                    if target_cust_type == "거주":
                        override_target['type'] = "거주"
                    else:
                        override_target['type'] = target_cust_type 
                        
                proposal = generate_marketing_text_with_gemini(
                    summary, 
                    persona, 
                    mbti_result, 
                    selected_mct, 
                    override_target=override_target if override_target else None
                )
                st.session_state['marketing_proposal'] = proposal
                st.session_state.chat_messages = []
                st.session_state.generating = False
                st.rerun()

        st.markdown("---")

        
        if st.session_state.get('marketing_proposal'):
            st.markdown(st.session_state['marketing_proposal'])
            
            st.markdown("---")
            st.subheader("🤖 AI 마케팅 도구 추천")
            st.info("아래 도구들과 프롬프트 예시를 활용하여 마케팅 콘텐츠를 손쉽게 제작해보세요.")

            reel_tab, blog_tab, image_tab = st.tabs(["🎬 **릴스/숏폼 제작**", "✍️ **블로그 포스팅**", "🎨 **이미지 생성**"])

            with reel_tab:
                st.markdown("""
                ### 🔹 Vrew  
                텍스트만 입력하면 자동으로 이미지, 영상 클립, 더빙까지 생성해주는 영상 제작 도구입니다. 릴스나 쇼츠 콘텐츠를 제작해 보세요!
                `https://vrew.voyagerx.com/`
                """)
                with st.expander("📝 **Vrew 활용 프롬프트 예시 펼쳐보기**"):
                    st.code(f"""
                    ### 릴스 대본 생성 프롬프트

                    **역할:**
                    당신은 '{summary['static_info'].get('HPSN_MCT_ZCD_NM')}' 가게를 운영하는 사장님 역할을 맡은 SNS 마케터입니다.
                    우리의 핵심 고객인 '{persona['name']}'의 관심을 끌 수 있는 30초 분량의 인스타그램 릴스 대본을 작성해주세요.

                    **릴스 컨셉:**
                    [사장님이 직접 가게의 매력을 소개하는 컨셉 / 고객이 직접 경험하는 듯한 1인칭 시점 컨셉 등]

                    **핵심 메시지:**
                    '{persona['goals'][0]}' 와 같은 고객의 니즈를 충족시키고, '{persona['pain_points'][0]}' 같은 불편함을 해결해준다는 점을 강조해주세요.

                    **포함할 내용:**
                    - 시선을 사로잡는 오프닝 멘트 (3초 이내)
                    - 가게의 핵심 메뉴 또는 서비스 소개
                    - 고객에게 제공하는 특별한 혜택 (이벤트, 할인 등)
                    - 행동 유도 문구 (예: "지금 바로 프로필 링크를 확인하세요!")
                    - 영상 장면에 대한 간단한 설명 (예: #1. 음식이 클로즈업되는 장면)

                    **분위기:**
                    [활기찬 / 감성적인 / 유머러스한] 분위기로 작성해주세요.
                    """, language="markdown")

            # ✍️ 블로그 탭
            with blog_tab:
                st.markdown("""
                ### 🔹 Gemini  
                강력한 AI 비서로 완성도 높은 블로그 글을 손쉽게 작성할 수 있어요!
                `https://gemini.google.com/`
                
                ### 🔹 뤼튼(Wrtn) 블로그  
                게시물의 주제, 말투를 설정하면 블로그 글을 자동으로 완성해 드려요!  
                `https://wrtn.ai/tools/67b2e7901b44a4d864b127a5`
                """)
                with st.expander("📝 **블로그 포스팅용 프롬프트 예시 펼쳐보기**"):
                    st.code(f"""
                    ### 블로그 포스트 생성 프롬프트

                    **역할:**
                    당신은 '{summary['static_info'].get('h_name', row.get('HPSN_MCT_BZN_CD_NM'))}' 상권의 맛집을 소개하는 전문 블로거입니다.

                    **주제:**
                    '{summary['static_info'].get('HPSN_MCT_ZCD_NM')}' 가게 방문 후기

                    **타겟 독자:**
                    '{persona['name']}' ({persona['description']})

                    **글의 목적:**
                    타겟 독자가 이 글을 읽고 우리 가게에 방문하고 싶게 만드는 것.
                    특히, '{persona['goals'][0]}'와 같은 독자의 목표를 우리 가게가 어떻게 만족시켜주는지 자연스럽게 녹여내 주세요.

                    **포함할 내용:**
                    1.  독자의 흥미를 유발하는 제목 (SEO 키워드: [지역명] 맛집, [업종명])
                    2.  가게의 첫인상 및 분위기 묘사
                    3.  주문한 메뉴와 맛에 대한 상세한 설명
                    4.  '{persona['pain_points'][0]}'과 같은 독자의 불편함을 우리 가게가 어떻게 해결해주는지에 대한 포인트 강조
                    5.  가게 위치, 운영 시간, 팁 등 방문 정보
                    6.  독자의 방문을 유도하는 마무리 문단

                    **글의 톤앤매너:**
                    [친근하고 솔직한 / 전문적이고 신뢰감 있는] 톤앤매너로 작성해주세요.
                    """, language="markdown")

            # 🎨 이미지 생성 탭
            with image_tab:
                st.markdown("""
                ### 🔹 뤼튼(Wrtn) 이미지  
                한국어에 강한 AI로, 손쉽게 원하는 이미지를 생성할 수 있어요! 
                `https://wrtn.ai/tools/67b2e7901b44a4d864b127b9`

                ### 🔹 Hailo AI  
                AI 에이전트를 활용하여 다양한 스타일의 이미지와 영상을 생성하고 편집할 수 있어요!
                `https://hailuoai.video/ko/agent`

                ### 🔹 플레이그라운드(로고) 
                간단한 입력만으로 원하는 이미지를 내 가게의 로고로 만들 수 있어요!  
                `https://playground.com/design/c/logo`
                """)

                with st.expander("📝 **이미지 생성 프롬프트 예시 펼쳐보기**"):
                    st.code(f"""
                    ### 마케팅 이미지 생성 프롬프트

                    **스타일:**
                    [실사 사진 / 디지털 아트 / 수채화 / 애니메이션 스타일]

                    **상세 설명:**
                    SNS 광고에 사용할 생동감 있고 매력적인 이미지.
                    '{summary['static_info'].get('HPSN_MCT_ZCD_NM')}' 식당에서 '{persona['name']}' 고객이 만족스럽게 식사를 즐기고 있는 장면.
                    '{persona['goals'][0]}'와 같은 기분을 느끼며 매우 만족스러워 보이는 표정.
                    분위기는 [아늑하고 따뜻한 / 밝고 현대적인 / 활기차고 트렌디한] 느낌.
                    메인 메뉴가 테이블 위에 아름답게 플레이팅 되어 있음.
                    디테일이 뛰어나고 따뜻하며 매력적인 조명에 초점을 맞출 것.

                    **핵심 키워드:**
                    맛있는 음식, 행복한 고객, {summary['static_info'].get('h_name', row.get('HPSN_MCT_BZN_CD_NM'))}, 라이프스타일, 고품질
                    """, language="markdown")

            st.markdown("---")
            st.subheader("🤖 추가 상담하기")
            
            if not st.session_state.chat_messages:
                st.session_state.chat_messages.append(
                    {"role": "assistant", "content": "방금 생성된 마케팅 전략에 대해 궁금한 점을 질문해보세요."}
                )

            for message in st.session_state.chat_messages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])

            if prompt := st.chat_input("여기에 질문을 입력하세요..."):
                st.session_state.chat_messages.append({"role": "user", "content": prompt})
                with st.chat_message("user"):
                    st.markdown(prompt)

                with st.chat_message("assistant"):
                    with st.spinner("AI가 답변을 생각 중입니다..."):
                        response = generate_chat_response_with_gemini(
                            base_context=f"분석정보: {summary}, 페르소나: {persona}, 원본전략: {st.session_state.marketing_proposal}",
                            messages_history=st.session_state.chat_messages
                        )
                        st.markdown(response)
                        st.session_state.chat_messages.append({"role": "assistant", "content": response})
            
            st.markdown("---") 
            
            docx_data = create_docx_report(
                selected_mct,
                st.session_state['marketing_proposal'],
                st.session_state.get('chat_messages', [])
            )
            
            st.download_button(
                label="📄 전체 내용 문서로 저장하기 (.docx)",
                data=docx_data,
                file_name=f"report_{selected_mct}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.info("👆 버튼을 눌러 우리 가게만을 위한 맞춤 마케팅 전략을 확인해보세요!")
    
if __name__ == '__main__':
    main()
